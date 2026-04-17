import sys
import subprocess
import os
import re

def install_and_import(package):
    try:
        __import__(package)
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])

install_and_import('docx')
from docx import Document
from docx.shared import Pt
import datetime

content = """# SportSight Analytics - Comprehensive MCA Project Report

## CHAPTER 1: INTRODUCTION

### 1.1 Introduction to the Project
The world of sports, particularly football (soccer), has become increasingly data-driven. However, sophisticated analytics platforms remain financially inaccessible to amateur clubs, local academies, and semi-professional teams. "SportSight Analytics" bridges this gap by providing an intuitive, web-based platform tailored for grassroots football. It allows coaches, team managers, and analysts to automate the ingestion of match day statistics and instantly translate raw numbers into actionable visual insights. By leveraging modern web frameworks and interactive charting libraries, the system eliminates tedious spreadsheet management, reduces human error, and delivers enterprise-grade analytical capabilities such as historical performance tracking, squad positional visualization, and direct one-on-one (1v1) player comparisons through an accessible interface.

### 1.2 Organization Profile
This project was designed, developed, and deployed to fulfill the major project requirements for the Master of Computer Applications (MCA) program. It demonstrates end-to-end full-stack software engineering proficiency, encompassing robust backend database architecture, secure user authentication systems, dynamic frontend data rendering, and comprehensive system testing. The project serves as a capstone, exhibiting practical application of academic concepts in database management systems, software engineering principles, UI/UX design, and complex algorithmic aggregations. 

### 1.3 Objectives of the Project
- To architect and build a centralized, secure web application for tracking and managing football match and player statistics.
- To automate data entry processes by developing a robust CSV ingestion engine that parses, validates, and stores bulk statistical data seamlessly.
- To improve data literacy and tactical decision-making for coaches by generating dynamic, real-time graphical representations of performance (e.g., rolling averages, interactive bar and line charts).
- To engineer a highly intuitive "1v1 Player Comparison" module that allows simultaneous evaluation of two distinct athletes against key performance indicators (KPIs).
- To implement strict role-based access control (RBAC), ensuring that only authorized administrative personnel can manipulate (create, update, delete) core statistical data while standard users maintain read-only analytical access.

### 1.4 Scope and Applicability
The current scope of SportSight Analytics predominantly covers association football (soccer). The system tracks domain-specific performance metrics such as goals, assists, tackles, saves, passing accuracy, and match ratings. It intelligently categorizes squad members into specific tactical roles (Attackers, Midfielders, Defenders, Goalkeepers) for grouped analysis. The application is highly applicable to local sporting academic institutions, university intramural leagues, and developmental football clubs seeking a low-cost, high-yield digital transformation of their performance assessment methodologies without the steep learning curves associated with professional enterprise tools.

## CHAPTER 2: REQUIREMENT AND ANALYSIS

### 2.1 Existing System / Problem Statement
Traditionally, local football clubs operate using highly un-optimized and fragmented data management systems. The standard practice involves scribbling statistics on physical paper ledgers during a match and later manually typing them into isolated desktop spreadsheet software (like MS Excel). 
The primary problems with this existing ecosystem are:
- High margin of human error during manual transcription.
- Severe lack of integration: data is decoupled from visualization tools, making complex queries (e.g., "Compare the last 5 match defensive ratings of our two center-backs") incredibly time-consuming.
- Difficulty in secure data sharing among coaching staff without creating multiple conflicting copies of spreadsheet files.
- Lack of immediate accessibility for players or fans to view their performance metrics.

### 2.2 Proposed System / Solution Overview
SportSight Analytics proposes a unified, cloud-ready centralized web platform tailored to resolve these exact issues. The solution provides a secure intake system where an Admin can upload a single, raw CSV file representing match-day stats. The backend server automatically sanitizes this data, resolves foreign key mappings to existing players and teams, and commits the records to a relational database. Immediately upon upload, the frontend analytics dashboard dynamically updates its graphical interfaces. Coaches can log in from any web-enabled device to view live, calculated rolling averages, filter squads by playing position, and instantly utilize the 1v1 comparison tool, turning a multi-hour analytical task into a single-click operation.

### 2.3 Feasibility Study
- Technical Feasibility: The project is highly feasible technically. Utilizing the Python-based Django framework provides a secure, "batteries-included" backend capable of rapid ORM (Object-Relational Mapping) database queries. The frontend relies on standard HTML5, CSS3, and JavaScript (Canvas API/Chart.js), ensuring cross-browser compatibility.
- Economic Feasibility: The application incurs virtually zero licensing costs as it leverages purely open-source technologies (Python, Django, SQLite/PostgreSQL, Chart.js, Bootstrap styling concepts). It can be hosted on very low-cost virtual private servers or platform-as-a-service (PaaS) providers.
- Operational Feasibility: The user interface follows modern, minimalistic design principles. Coaches, who may not be highly computer literate, do not require complex training modules; the system is self-explanatory, relying on familiar web patterns and straightforward navigation menus.

### 2.4 Conceptual Modelling
The conceptual architecture relies heavily on standard Entity-Relationship paradigms to ensure data integrity. The core entities include:
- Users (Admin, Staff, Standard) for authentication.
- Teams (representing the clubs).
- Players (linked via Foreign Key to a Team, storing inherent details like position and jersey number).
- Matches (representing a fixture on a specific date against an opponent).
- Performances (a transactional entity linking a Player and a Match, containing the specific metrics achieved during that 90-minute window).
This robust one-to-many relationship structure prevents data anomalies and simplifies aggregate querying for analytics.

### 2.5 Planning and Scheduling
The project lifecycle was strategically divided into agile sprints to ensure steady, measurable progress:
- Phase 1 (Week 1-2): Requirement gathering, conceptual schema design, and initial Django project initialization.
- Phase 2 (Week 3): Implementing User Authentication, Admin panels, and defining core database Models.
- Phase 3 (Week 4): Developing the complex algorithm for the bulk CSV data upload and validation pipeline.
- Phase 4 (Week 5): Developing Custom Views, routing, and integrating Chart.js for the dynamic Frontend Analytics Dashboard.
- Phase 5 (Week 6): Engineering the 1v1 comparison logic, resolving UI/UX styling issues (glassmorphism/dark mode adjustments), and comprehensive system testing and bug fixing.

## CHAPTER 3: SYSTEM SPECIFICATION

### 3.1 Software and Hardware Requirements
- Hardware Requirements:
  - Processor: Intel Core i3 / AMD Ryzen 3 (or equivalent) minimum.
  - Memory (RAM): Minimum 4 GB, Recommended 8 GB for smooth local development server execution.
  - Storage: 1 GB minimum free disk space.
- Software Requirements:
  - Operating System: Windows 10/11, macOS, or Linux distributions.
  - Runtime Environment: Python 3.10 or higher.
  - Backend Framework: Django 4.x / 5.x.
  - Database System: SQLite (Development) / PostgreSQL 14+ (Production via psycopg2).
  - Client Application: Any modern, HTML5-compliant web browser (Google Chrome, Mozilla Firefox, Microsoft Edge, Safari).

### 3.2 Functional Specifications
- Administrative Functionality: 
  - Exclusive access to secure administrative routes.
  - Ability to seamlessly upload bulk performance data via CSV files.
  - Complete CRUD (Create, Read, Update, Delete) privileges over Teams, Players, Matches, and discrete Performance records.
- Standard User Functionality:
  - Secure login and session management.
  - Access to the 'Season Analysis' dashboard featuring rich graphical visualizations (line charts, layered bar charts).
  - Ability to view team squads appropriately filtered and partitioned by their on-field roles (Attackers to Goalkeepers).
  - Full access to the 'Player Comparison' engine to analyze head-to-head metrics using visual progress bars and dynamic text styling.

### 3.3 Tools and Platforms Used
- Backend Core Programming: Python - chosen for its readability, rapid prototyping capabilities, and extensive standard library.
- Web Framework: Django - selected for its powerful ORM, built-in admin interface, CSRF protection, and MTV (Model-Template-View) architecture.
- Frontend Technologies: HTML5 for semantic structure, Vanilla CSS for responsive layout design, and Vanilla JavaScript for client-side interactivity and asynchronous rendering.
- Data Visualization Platform: Chart.js - chosen for its HTML5 canvas-based rendering, providing highly customizable, interactive, and responsive animated charts.
- Version Control & IDE: Git combined with GitHub for repository management, alongside Visual Studio Code (VS Code) acting as the primary Integrated Development Environment.

## CHAPTER 4: SYSTEM DESIGN

### 4.1 Module Descriptions
1. Authentication Module: Handles secure user registration, hashed password logins, session generation, and middleware protection to prevent unauthorized access to sensitive views (like data upload routes).
2. Data Ingestion Module: A backend Python module that safely reads uploaded CSV files, applies regex and type checking, gracefully strips unnecessary whitespaces, and instantiates database objects using Django's ORM operations.
3. Analytics Module: The core algorithmic engine residing in the views. It runs complex queries spanning multiple relational tables to compute overall team averages, individual player rolling averages, and aggregate positional metrics, serializing this into JSON.
4. Comparison Module: The 1v1 comparison engine that aggregates statistics for two selected players side-by-side using graphical progress bars.

### 4.2 Data / Schema Design
The backend schema is heavily normalized to guarantee data integrity:
- Table auth_user: Handles system credentials.
- Table core_team: Fields -> id, name, crest.
- Table core_player: Fields -> id, name, team_id, position, slug.
- Table core_match: Fields -> id, date, opponent.
- Table core_performance: Fields -> id, player_id, match_id, goals, assists, tackles, saves, pass_completion_rate.

### 4.3 Procedural / Flow Design
The system's control flow adheres to standard web request-response cycles:
- Admin Data Ingestion Flow: The Admin navigates to /upload, submits a file. The Django view catches the request in request.FILES, parses the CSV, validates relationships, saves objects, and redirects back.
- 1v1 Comparison Query Flow: The user selects two players. The backend runs aggregate queries to calculate average metrics for both Player A and Player B. The controller then dynamically calculates the mathematical difference and returns the result.

### 4.4 User Interface Design
The user interface breaks away from traditional administrative panels by employing a highly modern, "glassmorphism" aesthetic. The design relies on a deep, dark-mode background accented with vibrant colors tailored to functional context (e.g., offensive stats in blue, defensive in green). The UI uses CSS Grid and Flexbox to ensure full responsiveness across mobile and desktop interfaces.

## CHAPTER 5: AGILE / DEVELOPMENT METHODOLOGY

### 5.1 Project Roadmap / Schedule
Development adhered strictly to Agile methodologies, iterating over short sprints. After initial structural mockups and schema definitions, the first MVP was simply the ability to add a player and record a stat. Subsequent sprints layered complexity: First adding the automated CSV upload, then implementing complex ORM aggregations to feed Chart.js, and finally polishing the complex UI interactions.

### 5.2 User Stories / Sprint Planning
User stories drove feature implementation:
- "As an Admin, I require the ability to rapidly ingest multi-row CSVs containing 20+ players' match details without creating duplicate entries."
- "As an Analyst, I need an immediate visual representation combining a player's raw goals and their overall performance trendline over the season."
- "As a standard User, I need to easily access a dedicated page to compare my favorite striker against our main defender."

### 5.3 Test Plan
Testing encompassed functional, integration, and UI verification. Unit tests evaluated backend data parsing handling highly corrupted CSV files. Integration testing ensured database consistency, specifically verifying that Performance record creations successfully trigger analytics recalculations. UI tests focused on chart canvas rendering and responsiveness.

## CHAPTER 6: IMPLEMENTATION AND TESTING

### 6.1 Implementation Procedures
Implementation rigidly followed the Django MVT architecture. Models were defined in Python and migrated to SQL via makemigrations. Complex business logic, such as data processing and statistical mathematics, was kept out of templates and encapsulated within views.py. The frontend simply iterates over processed context dictionaries using the Django Template Language.

### 6.2 Testing Methods and Results
During implementation, several edge cases were identified and resolved. A specific TemplateSyntaxError involving the slugify filter in the comparison engine was resolved by moving the logic to the backend. Data integrity bugs post-CSV uploads (e.g., players erroneously changing teams) were fixed by enhancing database constraints. The system has passed rigorous evaluation.

## CHAPTER 7: CONCLUSIONS

### 7.1 Summary
The "SportSight Analytics" project successfully delivered a robust, highly optimized, and visually engaging web application tailored for real-world sports data management. It functionally replaces tedious manual documentation processes with an automated system capable of bulk ingestions and immediate, complex visual readouts. 

### 7.2 Limitations
In its current state, the system relies on post-match, manual upload scenarios and does not yet ingest live telemetry feed metrics. The statistical schemas are heavily hardcoded to track football-specific KPIs, meaning adaptation to other sports requires fundamental database restructuring.

### 7.3 Future Scope
Future iterations hold immense potential. The architecture is primed for integration with Machine Learning models (e.g., Scikit-Learn) to build predictive models for injury risks based on match load. Additionally, building REST API consumers to automatically ingest data from external sources like Sportmonks would further automate the analysis pipeline.

## CHAPTER 8: APPENDICES

### 8.1 Screenshots
(Attach the following screenshots here in the final Word Document)
- Login Screen / Registration Gateway
- Dynamic Season Analysis Dashboard (Charts page)
- Positional Squad Layouts Page
- 1v1 Player Comparison Results
- Administrative CSV Data Upload Form

### 8.2 Sample Code / Queries
The project contains complex algorithms, specifically in the CSV parser within views.py and the Chart.js integration scripts within themes.js. These manage the asynchronous generation of interactive bar/line composites and should be included for reference.

## CHAPTER 9: REFERENCES

1. Django Project Foundation, Official Django Documentation, djangoproject.com
2. Chart.js Contributors, Chart.js Data Visualization API, chartjs.org
3. Mozilla Developer Network (MDN), HTML5 Canvas API and JavaScript, developer.mozilla.org
4. Python Software Foundation, Standard Library Documentation, docs.python.org
"""

doc = Document()

for line in content.split('\\n'):
    line = line.strip()
    if not line:
        continue
    
    line_clean = line.replace('**', '').replace('*', '')

    if line.startswith('# '):
        doc.add_heading(line_clean.replace('# ', '', 1), level=1)
    elif line.startswith('## '):
        doc.add_heading(line_clean.replace('## ', '', 1), level=2)
    elif line.startswith('### '):
        doc.add_heading(line_clean.replace('### ', '', 1), level=3)
    elif line.startswith('- '):
        doc.add_paragraph(line_clean.replace('- ', '', 1), style='List Bullet')
    elif re.match(r'^\d+\.', line):
        doc.add_paragraph(line_clean, style='List Paragraph')
    elif line.startswith('---'):
        pass
    else:
        doc.add_paragraph(line_clean)

# Append dynamic generation timestamp
doc.add_paragraph(f"\\nDocument Generated On: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

output_path = r"c:\Users\alber\OneDrive\Desktop\all filess\SUIII\MCA PROJECT BACKUP\sportsight_backup\SportSight_Complete_Report_Detailed.docx"
doc.save(output_path)
print(f"Successfully generated full detailed report at {output_path}")
