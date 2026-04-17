import sys
import os
import subprocess

def install_and_import(package):
    try:
        __import__(package)
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])

install_and_import('docx')
from docx import Document

doc = Document()
doc.add_heading('SportSight Analytics - Database Tables', level=1)
doc.add_paragraph('These tables represent the Database Schema for your MCA project report.')

tables_info = {
    "1. Users Collection Schema": [
        ("Field Name", "Data Type", "Description"),
        ("id", "Integer", "Primary Key, Auto-increment"),
        ("username", "String", "Unique login username"),
        ("password", "String", "Hashed password"),
        ("is_staff", "Boolean", "True if Admin, False if Standard User"),
        ("email", "String", "User's email address")
    ],
    "2. Teams Collection Schema": [
        ("Field Name", "Data Type", "Description"),
        ("id", "Integer", "Primary Key, Auto-increment"),
        ("name", "String", "Name of the team (e.g., Liverpool)"),
        ("logo", "ImageField", "Path to the team's logo image"),
        ("created_at", "DateTime", "Timestamp of creation")
    ],
    "3. Players Collection Schema": [
        ("Field Name", "Data Type", "Description"),
        ("id", "Integer", "Primary Key, Auto-increment"),
        ("name", "String", "Full name of the player"),
        ("position", "String", "Attacker, Midfielder, Defender, or Goalkeeper"),
        ("team_id", "Integer", "Foreign Key referencing Teams table"),
        ("jersey_number", "Integer", "Player's uniform number")
    ],
    "4. Matches Collection Schema": [
        ("Field Name", "Data Type", "Description"),
        ("id", "Integer", "Primary Key, Auto-increment"),
        ("date", "Date", "Date the match was played"),
        ("opponent", "String", "Name of the opposing team"),
        ("result", "String", "Win, Loss, or Draw")
    ],
    "5. Performances Collection Schema": [
        ("Field Name", "Data Type", "Description"),
        ("id", "Integer", "Primary Key, Auto-increment"),
        ("player_id", "Integer", "Foreign Key referencing Players table"),
        ("match_id", "Integer", "Foreign Key referencing Matches table"),
        ("goals", "Integer", "Number of goals scored in the match"),
        ("assists", "Integer", "Number of assists provided"),
        ("tackles", "Integer", "Number of successful tackles"),
        ("saves", "Integer", "Number of saves (Goalkeepers)"),
        ("passes_completed", "Integer", "Total successful passes"),
        ("rating", "Float", "Overall match rating out of 10.0")
    ]
}

for title, rows in tables_info.items():
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Add Header Row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = rows[0][0]
    hdr_cells[1].text = rows[0][1]
    hdr_cells[2].text = rows[0][2]
    
    # Add Data Rows
    for row_data in rows[1:]:
        row_cells = table.add_row().cells
        row_cells[0].text = row_data[0]
        row_cells[1].text = row_data[1]
        row_cells[2].text = row_data[2]
        
    doc.add_paragraph('\n')

output_path = r"c:\Users\alber\OneDrive\Desktop\all filess\SUIII\MCA PROJECT BACKUP\sportsight_backup\SportSight_Tables.docx"
doc.save(output_path)
print(f"Successfully generated {output_path}")
